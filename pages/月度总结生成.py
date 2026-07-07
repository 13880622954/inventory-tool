import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from datetime import datetime, timedelta
from openpyxl import load_workbook
import re

# ---------------------------- 辅助函数 ----------------------------
def excel_date_to_str(serial):
    if isinstance(serial, (int, float)) and serial > 40000:
        dt = datetime(1899, 12, 30) + timedelta(days=serial)
        return f"{dt.year}年{dt.month}月"
    return str(serial)

def parse_number(text):
    match = re.search(r'(\d+)', str(text))
    return int(match.group(1)) if match else 0

def get_current_date():
    return datetime.now().strftime("%Y年%m月%d日")

def get_data_month(current_date):
    y, m = current_date.year, current_date.month - 1
    if m == 0:
        m = 12
        y -= 1
    return y, m

def get_next_month(current_date):
    y, m = current_date.year, current_date.month
    if m == 12:
        return y + 1, 1
    else:
        return y, m + 1

def get_current_month_first_monday(year, month):
    first_day = datetime(year, month, 1)
    start = first_day
    while start.weekday() != 0:
        start += timedelta(days=1)
    return start

# ---------------------------- 读取 Excel 数据 ----------------------------
def load_excel_data(excel_file):
    wb = load_workbook(excel_file, data_only=True)
    data = {}

    # 基本信息
    if '基本信息' in wb.sheetnames:
        ws = wb['基本信息']
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1] is not None:
                data[str(row[0])] = row[1]
    raw_month = data.get('报告月份', '')
    data['报告月份'] = excel_date_to_str(raw_month) if isinstance(raw_month, (int, float)) else str(raw_month)
    data['上月I030预留未清'] = int(data.get('上月I030预留未清', 0))
    data['上月I030销售未清'] = int(data.get('上月I030销售未清', 0))
    data['上月IB00预留未清'] = int(data.get('上月IB00预留未清', 0))
    data['上月IB00销售未清'] = int(data.get('上月IB00销售未清', 0))

    # 对账差异
    df_diff = pd.read_excel(excel_file, sheet_name='对账差异', header=0)
    row = df_diff[df_diff['产品线'] == '美菱']
    if not row.empty:
        diff_text = str(row.iloc[0, 1])
        data['diff_total'] = parse_number(diff_text)
        data['diff_handled'] = parse_number(re.search(r'已处理(\d+)台', diff_text).group(1)) if re.search(r'已处理(\d+)台', diff_text) else 0
        data['diff_unhandled'] = parse_number(re.search(r'未处理(\d+)台', diff_text).group(1)) if re.search(r'未处理(\d+)台', diff_text) else 0
        data['diff_cancel'] = row.iloc[0, 2] if len(row.columns) > 2 else ''
        data['diff_empty_order'] = row.iloc[0, 3] if len(row.columns) > 3 else ''
        data['diff_reason'] = row.iloc[0, 4] if len(row.columns) > 4 else ''

    # 对账差异明细（异常差异）
    data['abnormal_list'] = []
    df_detail = pd.read_excel(excel_file, sheet_name='对账差异明细')
    df_detail = df_detail[df_detail.iloc[:, 0] != '总计']
    abnormal = df_detail[df_detail['异常差异'] > 0]
    data['abnormal_list'] = abnormal.iloc[:, 0].tolist()

    # 只读取盘点成品（不包含赠品）
    total_books = total_actual = total_in = total_out = 0
    substantive_diff = 0
    if '盘点成品' in wb.sheetnames:
        ws = wb['盘点成品']
        for row in ws.iter_rows(min_row=3, max_row=ws.max_row, values_only=True):
            if row[0] == '合计':
                break
            if row[1] and isinstance(row[1], (int, float)):
                total_books += row[1]
                total_actual += (row[6] or 0)
                total_in += (row[2] or 0)
                total_out += (row[3] or 0)
                substantive_diff += (row[4] or 0) + (row[5] or 0)
    data['total_books'] = total_books
    data['total_actual'] = total_actual
    data['pan_ying'] = total_in
    data['pan_kui'] = total_out
    data['substantive_diff'] = "无" if substantive_diff == 0 else f"盘盈{total_in}台，盘亏{total_out}台"
    data['diff_reason_stock'] = "1. C端订单因信用、价格问题未维护，客户代码冻结导致R3下账失败，其他异常原因导致R3下账失败"

    # 当前日期和数据月份
    current_date = datetime.now()
    data_month = get_data_month(current_date)[1]
    data['data_month'] = data_month
    data['current_month'] = current_date.month
    data['report_month'] = f"{current_date.year}年{get_data_month(current_date)[1]}月"
    data['current_date'] = get_current_date()

    # I030未清
    i030_reserve_current = 0
    i030_sale_current = 0
    if 'I030销售预留未清及处理情况' in wb.sheetnames:
        ws = wb['I030销售预留未清及处理情况']
        for row in ws.iter_rows(values_only=True):
            if row[0] == '合计':
                i030_reserve_current = row[3] or 0
                i030_sale_current = row[7] or 0
                break
    data['i030_reserve_current'] = i030_reserve_current
    data['i030_sale_current'] = i030_sale_current

    # I030 销售未清
    if data['上月I030销售未清'] > i030_sale_current:
        diff = data['上月I030销售未清'] - i030_sale_current
        data['i030_sale_processed'] = diff
        data['i030_sale_note'] = f"{data_month}月已处理{diff}单"
    elif data['上月I030销售未清'] < i030_sale_current:
        diff = i030_sale_current - data['上月I030销售未清']
        data['i030_sale_processed'] = 0
        data['i030_sale_note'] = f"增加{diff}个"
    else:
        data['i030_sale_processed'] = 0
        data['i030_sale_note'] = ""

    # I030 预留未清
    if data['上月I030预留未清'] > i030_reserve_current:
        diff = data['上月I030预留未清'] - i030_reserve_current
        data['i030_reserve_processed'] = diff
        data['i030_reserve_note'] = f"{data_month}月已处理{diff}单"
    elif data['上月I030预留未清'] < i030_reserve_current:
        diff = i030_reserve_current - data['上月I030预留未清']
        data['i030_reserve_processed'] = 0
        data['i030_reserve_note'] = f"增加{diff}个"
    else:
        data['i030_reserve_processed'] = 0
        data['i030_reserve_note'] = ""

    # IB00未清
    ib00_reserve_current = 0
    ib00_sale_current = 0
    if 'IB00销售、预留未清订单' in wb.sheetnames:
        ws = wb['IB00销售、预留未清订单']
        for row in ws.iter_rows(values_only=True):
            if row[0] == '合计':
                ib00_reserve_current = row[1] or 0
                ib00_sale_current = row[2] or 0
                break
    data['ib00_reserve_current'] = ib00_reserve_current
    data['ib00_sale_current'] = ib00_sale_current

    # IB00 销售未清（保留原固定备注，追加处理信息）
    base_note = "关于历史订单的跟进情况：太原2单、重庆1单石家庄1单已反馈给货源经理，反馈无法查到USO单号无法删除,目前暂时无法处理，2025年的订单为2单为OFC切换到EBOC系统时未清理现在已无法清理，"
    if data['上月IB00销售未清'] > ib00_sale_current:
        diff = data['上月IB00销售未清'] - ib00_sale_current
        data['ib00_sale_processed'] = diff
        data['ib00_sale_note'] = base_note + f"{data_month}月已处理{diff}单"
    elif data['上月IB00销售未清'] < ib00_sale_current:
        diff = ib00_sale_current - data['上月IB00销售未清']
        data['ib00_sale_processed'] = 0
        data['ib00_sale_note'] = base_note + f"增加{diff}个"
    else:
        data['ib00_sale_processed'] = 0
        data['ib00_sale_note'] = base_note

    # IB00 预留未清
    if data['上月IB00预留未清'] > ib00_reserve_current:
        diff = data['上月IB00预留未清'] - ib00_reserve_current
        data['ib00_reserve_processed'] = diff
        data['ib00_reserve_note'] = f"{data_month}月已处理{diff}单"
    elif data['上月IB00预留未清'] < ib00_reserve_current:
        diff = ib00_reserve_current - data['上月IB00预留未清']
        data['ib00_reserve_processed'] = 0
        data['ib00_reserve_note'] = f"增加{diff}个"
    else:
        data['ib00_reserve_processed'] = 0
        data['ib00_reserve_note'] = ""

    # 借用库箱机损库（超期60天）- 按库存地尾数汇总
    overdue = {'干线箱损库': 0, '配送箱损库': 0, '配送运损库': 0, '借用库': 0, '干线运损库': 0}
    if '借用库箱机损库（超期60天）' in wb.sheetnames:
        ws = wb['借用库箱机损库（超期60天）']
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) < 4:
                continue
            kucun_di = row[1]
            value_2m_above = row[3]
            if not kucun_di or not isinstance(value_2m_above, (int, float)):
                continue
            last_digit = None
            if kucun_di and str(kucun_di)[-1].isdigit():
                last_digit = int(str(kucun_di)[-1])
            if last_digit == 4:
                overdue['干线箱损库'] += value_2m_above
            elif last_digit == 5:
                overdue['配送箱损库'] += value_2m_above
            elif last_digit == 8:
                overdue['借用库'] += value_2m_above
            elif last_digit == 3:
                overdue['配送运损库'] += value_2m_above
            elif last_digit == 2:
                overdue['干线运损库'] += value_2m_above
    data['overdue'] = overdue

    # 批扫
    left_in = left_out = right_in = right_out = 0
    right_regions = set()
    if '批扫' in wb.sheetnames:
        ws = wb['批扫']
        for row in ws.iter_rows(min_row=4, max_col=5, values_only=True):
            if row[0] and not any(kw in str(row[0]) for kw in ['转库批扫', '产品线', '出库', '入库', '物料代码']):
                if isinstance(row[2], (int, float)):
                    left_in += row[2]
                    left_out += row[3] or 0
        for row in ws.iter_rows(min_row=4, min_col=6, max_col=10, values_only=True):
            if row[0] and not any(kw in str(row[0]) for kw in ['非转库批扫', '产品线', '数量', '涉及', '物料代码']):
                if isinstance(row[2], (int, float)):
                    right_in += row[2]
                    right_out += row[3] or 0
                    if row[1]:
                        right_regions.add(str(row[1]))
    data['batch_left_out'] = left_out
    data['batch_left_in'] = left_in
    data['batch_right_out'] = right_out
    data['batch_right_in'] = right_in
    data['batch_regions'] = '、'.join(right_regions) if right_regions else '无'
    data['batch_violation_regions'] = '无'

    # WMS未清
    wms_tai = 0
    wms_dan = 0
    if 'WMS未清' in wb.sheetnames:
        df_wms = pd.read_excel(excel_file, sheet_name='WMS未清')
        dan_col = None
        tai_col = None
        for col in df_wms.columns:
            if '出库单号' in col:
                dan_col = col
            if '机号码数量' in col:
                tai_col = col
        if dan_col:
            wms_dan = df_wms[dan_col].nunique()
        if tai_col:
            wms_tai = df_wms[tai_col].sum()
    data['wms_tai'] = wms_tai
    data['wms_dan'] = wms_dan
    data['wms_quantity'] = f"{wms_tai}台、{wms_dan}单"

    # ================== 纸质盘点表总结（修正动态月份） ==================
    if '纸质盘点表及扫描件回传' in wb.sheetnames:
        df_paper = pd.read_excel(excel_file, sheet_name='纸质盘点表及扫描件回传')
        total = len(df_paper)

        # 动态获取月份列名：如 "5月"、"6月" 等
        month_col = f"{data['data_month']}月"
        # 如果列名不存在，自动查找含“月”的列（兼容命名差异）
        if month_col not in df_paper.columns:
            candidates = [col for col in df_paper.columns if '月' in str(col)]
            month_col = candidates[0] if candidates else '4月'  # 兜底

        paper_rec = df_paper[df_paper[month_col] == '已回'].shape[0]
        scan_rec = df_paper[df_paper['扫描件回收情况'] == '已回'].shape[0]
        missing_paper = df_paper[df_paper[month_col] != '已回']['地区（美菱）'].tolist()
        missing_scan = df_paper[df_paper['扫描件回收情况'] != '已回']['地区（美菱）'].tolist()
        anomalies = []
        for _, row in df_paper.iterrows():
            if pd.notna(row['其他备注']) or pd.notna(row['回寄情况']):
                anomalies.append(f"{row['地区（美菱）']}: {row['其他备注'] or ''} {row['回寄情况'] or ''}".strip())
        summary = f"共{total}个仓库，纸质盘点表已回收{paper_rec}个，缺失{len(missing_paper)}个：" + ('、'.join(missing_paper) if missing_paper else '无')
        summary += f"；扫描件已回收{scan_rec}个，缺失{len(missing_scan)}个：" + ('、'.join(missing_scan) if missing_scan else '无')
        data['paper_check_summary'] = summary
        data['paper_remarks'] = '；'.join(anomalies) if anomalies else '无'
    else:
        data['paper_check_summary'] = "纸质盘点表核对情况：无数据"
        data['paper_remarks'] = "无"

    # 机损台账数据准备
    jisun_rows = []
    if '机损' in wb.sheetnames:
        ws = wb['机损']
        rows = list(ws.iter_rows(values_only=True))
        last_start = current_start = None
        for i, row in enumerate(rows):
            if row and row[0] == '上月机损数据':
                last_start = i+2
            if row and row[0] == '本月机损数据':
                current_start = i+2
        last_dict = {}
        if last_start:
            for i in range(last_start, len(rows)):
                row = rows[i]
                if not row[0] or row[0] == '' or row[0] == '本月机损数据':
                    break
                if row[0] and len(row) >= 6:
                    code = str(row[2]) if row[2] else ''
                    last_dict[code] = {
                        '库位': row[0] or '',
                        '库位描述': row[1] or '',
                        '物料代码': code or '',
                        '批次': row[3] or '',
                        '库存': row[4] if row[4] is not None else 0,
                        '进度': row[5] or '',
                        '新增库存': row[6] if len(row) > 6 and row[6] is not None else ''
                    }
        current_dict = {}
        if current_start:
            for i in range(current_start, len(rows)):
                row = rows[i]
                if not row[0] or row[0] == '':
                    break
                if row[0] and len(row) >= 5:
                    code = str(row[0])
                    batch = row[6] if len(row) > 6 else ''
                    current_dict[code] = {
                        '库位': row[3] if len(row) > 3 and row[3] is not None else '',
                        '库位描述': row[5] if len(row) > 5 and row[5] is not None else '',
                        '物料代码': code or '',
                        '批次': batch or '',
                        '库存': row[2] if len(row) > 2 and row[2] is not None else 0,
                        '进度': '',
                        '新增库存': ''
                    }
        all_codes = set(last_dict.keys()) | set(current_dict.keys())
        for code in all_codes:
            if code in last_dict and code in current_dict:
                item = last_dict[code]
                item['进度'] = '重点关注'
                item['新增库存'] = '在周期内'
                jisun_rows.append(item)
            elif code in last_dict and code not in current_dict:
                item = last_dict[code]
                item['进度'] = '已处理' if item['进度'] else '已处理'
                item['新增库存'] = ''
                jisun_rows.append(item)
            elif code not in last_dict and code in current_dict:
                item = current_dict[code]
                item['进度'] = '新增'
                item['新增库存'] = '在周期内'
                jisun_rows.append(item)
    data['jisun_rows'] = jisun_rows

    # 主要问题分析数据
    issues_rows = []
    for w in data['abnormal_list']:
        issues_rows.append({'事项分类': '其他问题', '所属问题类型': '对账差异', '问题描述分析': f'{w}存在异常差异，需跟进处理。', '解决措施': '已联系对应人员核实解决。'})
    if data['overdue']['干线运损库'] > 0:
        issues_rows.append({'事项分类': '库存积压', '所属问题类型': '超期60天', '问题描述分析': f"干线运损库超期{data['overdue']['干线运损库']}台，需关注。", '解决措施': '分析原因，加快处理进度。'})
    data['issues_rows'] = issues_rows

    # 其他占位符
    data['abnormal_text'] = '；'.join([f"涉及区域{w}，因异常问题造成差异，处理情况：正在跟进处理。" for w in data['abnormal_list']]) or '无'
    data['create_order_num'] = '【请填写】'
    data['complement_num'] = '【请填写】'

    return data

# ---------------------------- 更新 Word 文档 ----------------------------
def update_word_document(template_bytes, data):
    doc = Document(BytesIO(template_bytes))
    current_date = datetime.now()
    data_month = data['data_month']
    current_month = data['current_month']

    # 普通占位符替换字典
    replacements = {
        '{{report_month}}': data['report_month'],
        '{{current_date}}': data['current_date'],
        '{{abnormal_text}}': data['abnormal_text'],
        '{{diff_total}}': str(data['diff_total']),
        '{{diff_handled}}': str(data['diff_handled']),
        '{{diff_unhandled}}': str(data['diff_unhandled']),
        '{{diff_empty_order}}': data['diff_empty_order'],
        '{{diff_cancel}}': data['diff_cancel'],
        '{{diff_reason}}': data['diff_reason'],
        '{{wms_quantity}}': data['wms_quantity'],
        '{{data_month}}': str(data_month),
        '{{I030_sale_rdc}}': str(data['i030_sale_current']),
        '{{I030_sale_nation}}': "",
        '{{I030_sale_processed}}': str(data['i030_sale_processed']),
        '{{I030_sale_note}}': data['i030_sale_note'],
        '{{I030_reserve_rdc}}': str(data['i030_reserve_current']),
        '{{I030_reserve_nation}}': "",
        '{{I030_reserve_processed}}': str(data['i030_reserve_processed']),
        '{{I030_reserve_note}}': data['i030_reserve_note'],
        '{{IB00_sale_quantity}}': str(data['ib00_sale_current']),
        '{{IB00_sale_processed}}': str(data['ib00_sale_processed']),
        '{{IB00_sale_note}}': data['ib00_sale_note'],
        '{{IB00_reserve_quantity}}': str(data['ib00_reserve_current']),
        '{{IB00_reserve_processed}}': str(data['ib00_reserve_processed']),
        '{{IB00_reserve_note}}': data['ib00_reserve_note'],
        '{{i030_reserve_current}}': str(data['i030_reserve_current']),
        '{{ib00_reserve_processed}}': str(data['ib00_reserve_processed']),
        '{{batch_left_out}}': str(data['batch_left_out']),
        '{{batch_left_in}}': str(data['batch_left_in']),
        '{{batch_right_out}}': str(data['batch_right_out']),
        '{{batch_regions}}': data['batch_regions'],
        '{{current_month}}': str(current_month),
        '{{batch_violation_regions}}': data['batch_violation_regions'],
        '{{total_books}}': str(data['total_books']),
        '{{total_actual}}': str(data['total_actual']),
        '{{pan_ying}}': str(data['pan_ying']),
        '{{pan_kui}}': str(data['pan_kui']),
        '{{substantive_diff}}': data['substantive_diff'],
        '{{diff_reason_stock}}': data['diff_reason_stock'],
        '{{paper_check_summary}}': data['paper_check_summary'],
        '{{paper_remarks}}': data['paper_remarks'],
        '{{create_order_num}}': data['create_order_num'],
        '{{complement_num}}': data['complement_num'],
        '{{overdue_ganxianxiang}}': str(data['overdue']['干线箱损库']),
        '{{overdue_peisongxiang}}': str(data['overdue']['配送箱损库']),
        '{{overdue_peisongyun}}': str(data['overdue']['配送运损库']),
        '{{overdue_jieyong}}': str(data['overdue']['借用库']),
        '{{overdue_ganxianyun}}': str(data['overdue']['干线运损库']),
    }

    # 替换段落中的占位符
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in para.text:
                            para.text = para.text.replace(key, value)

    # 处理机损台账
    for table in doc.tables:
        rows_to_remove = []
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if '{{jisun_rows}}' in cell.text:
                    rows_to_remove.append(row)
                    break
        for row in rows_to_remove:
            tbl = row._element.getparent()
            tbl.remove(row._element)
            for item in data['jisun_rows']:
                new_row = table.add_row()
                if len(new_row.cells) >= 7:
                    new_row.cells[0].text = str(item.get('库位', ''))
                    new_row.cells[1].text = str(item.get('库位描述', ''))
                    new_row.cells[2].text = str(item.get('物料代码', ''))
                    new_row.cells[3].text = str(item.get('批次', ''))
                    new_row.cells[4].text = str(item.get('库存', ''))
                    new_row.cells[5].text = str(item.get('进度', ''))
                    new_row.cells[6].text = str(item.get('新增库存', ''))
            break

    # 处理主要问题分析表格
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if '{{issues_table}}' in cell.text:
                    tbl = row._element.getparent()
                    tbl.remove(row._element)
                    for issue in data['issues_rows']:
                        new_row = table.add_row()
                        if len(new_row.cells) >= 4:
                            new_row.cells[0].text = str(issue.get('事项分类', ''))
                            new_row.cells[1].text = str(issue.get('所属问题类型', ''))
                            new_row.cells[2].text = str(issue.get('问题描述分析', ''))
                            new_row.cells[3].text = str(issue.get('解决措施', ''))
                    break
            else:
                continue
            break

    # 下月工作计划
    year, month = current_date.year, current_date.month
    first_monday = get_current_month_first_monday(year, month)
    weeks = [(first_monday + timedelta(days=7*i), first_monday + timedelta(days=7*i+6)) for i in range(4)]
    week_names = ['第一周', '第二周', '第三周', '第四周']
    next_month = get_next_month(current_date)[1]

    for para in doc.paragraphs:
        for i, week_name in enumerate(week_names):
            if week_name in para.text:
                start, end = weeks[i]
                date_range = f"（{start.strftime('%m月%d日')}-{end.strftime('%m月%d日')}）"
                if date_range not in para.text:
                    para.text = re.sub(r'（\d+月\d+日-\d+月\d+日）', '', para.text)
                    para.text = para.text.replace(week_name, f"{week_name}{date_range}")
                break
        if '将' in para.text and '月盘点表' in para.text:
            para.text = re.sub(r'将\d+月盘点表', f'将{data_month}月盘点表', para.text)
        if '核对' in para.text and '月WMS批扫数据' in para.text:
            para.text = re.sub(r'核对\d+月WMS批扫数据', f'核对{data_month}月WMS批扫数据', para.text)
        if '核对、回收' in para.text and '月美菱纸质盘点表' in para.text:
            para.text = re.sub(r'核对、回收\d+月美菱纸质盘点表', f'核对、回收{data_month}月美菱纸质盘点表', para.text)
        if '为' in para.text and '月盘存做准备' in para.text:
            para.text = re.sub(r'为\d+月盘存做准备', f'为{next_month}月盘存做准备', para.text)

    # 特殊：整改措施中的数字追加（销售+预留总和）
    i030_total_processed = data['i030_sale_processed'] + data['i030_reserve_processed']
    ib00_total_processed = data['ib00_sale_processed'] + data['ib00_reserve_processed']

    for para in doc.paragraphs:
        # 第一条整改措施：I030 总处理数量
        if '余下' in para.text and '单为异常订单暂不处理' in para.text:
            if i030_total_processed > 0:
                if f"{data_month}月已处理" not in para.text:
                    para.text = para.text + f"，{data_month}月已处理{i030_total_processed}单"
        # 第二条整改措施：IB00 总处理数量
        if 'IB00工厂2025年之前的遗留预留订单未清项' in para.text and '月处理' in para.text:
            new_text = re.sub(r'\d+月处理\d+单', f"{data_month}月处理{ib00_total_processed}单", para.text)
            if new_text != para.text:
                para.text = new_text

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ---------------------------- Streamlit UI ----------------------------
def main():
    st.set_page_config(page_title="月度总结自动生成器", layout="centered")
    st.title("📊 数据组月度总结自动生成器")
    st.markdown("上传 Excel 数据文件和 Word 模板（已添加占位符），自动生成更新后的月度总结报告。")

    uploaded_excel = st.file_uploader("上传 Excel 数据文件", type=["xlsx"])
    uploaded_word = st.file_uploader("上传 Word 模板文件（带占位符）", type=["docx"])

    if st.button("生成报告"):
        if uploaded_excel and uploaded_word:
            try:
                data = load_excel_data(uploaded_excel)
                new_word_bytes = update_word_document(uploaded_word.read(), data)
                st.success("生成成功！")
                st.download_button(
                    label="下载 Word 报告",
                    data=new_word_bytes,
                    file_name=f"数据组月度工作总结及下月计划（{data['report_month']}）.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"生成失败：{str(e)}")
                st.exception(e)
        else:
            st.warning("请上传两个文件")

if __name__ == "__main__":
    main()
